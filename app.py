from __future__ import annotations

import io
import os
from pathlib import Path

import pdfplumber
import pytesseract
import streamlit as st
from adobe.pdfservices.operation.auth.service_principal_credentials import ServicePrincipalCredentials
from adobe.pdfservices.operation.exception.exceptions import ServiceApiException, ServiceUsageException, SdkException
from adobe.pdfservices.operation.io.cloud_asset import CloudAsset
from adobe.pdfservices.operation.io.stream_asset import StreamAsset
from adobe.pdfservices.operation.pdf_services import PDFServices
from adobe.pdfservices.operation.pdf_services_media_type import PDFServicesMediaType
from adobe.pdfservices.operation.pdfjobs.jobs.export_pdf_job import ExportPDFJob
from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_params import ExportPDFParams
from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_target_format import ExportPDFTargetFormat
from adobe.pdfservices.operation.pdfjobs.result.export_pdf_result import ExportPDFResult
from docx import Document
from docx.shared import Inches, Pt
from openpyxl import Workbook
from pdf2image import convert_from_bytes
from PIL import Image
from pytesseract import Output


OUTPUT_FORMATS = {
    "Word (.docx)": {
        "extension": "docx",
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "target_format": ExportPDFTargetFormat.DOCX,
    },
    "Excel (.xlsx)": {
        "extension": "xlsx",
        "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "target_format": ExportPDFTargetFormat.XLSX,
    },
    "PowerPoint (.pptx)": {
        "extension": "pptx",
        "mime": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        "target_format": ExportPDFTargetFormat.PPTX,
    },
}


def set_page_style() -> None:
    st.set_page_config(page_title="PDF変換アプリ", page_icon="📄", layout="centered")
    st.markdown(
        """
        <style>
            .stApp {
                background: linear-gradient(180deg, #f7f9fc 0%, #eef3f9 100%);
            }
            .block-container {
                max-width: 900px;
                padding-top: 2rem;
                padding-bottom: 2rem;
            }
            .hero {
                background: white;
                border: 1px solid rgba(15, 23, 42, 0.08);
                border-radius: 20px;
                padding: 1.5rem 1.6rem;
                box-shadow: 0 16px 40px rgba(15, 23, 42, 0.08);
                margin-bottom: 1.25rem;
            }
            .hero h1 {
                margin: 0;
                font-size: 1.8rem;
            }
            .hero p {
                margin: 0.45rem 0 0;
                color: #475569;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )


def read_config_value(key: str) -> str:
    env_value = os.getenv(key, "")
    if env_value:
        return env_value

    try:
        secret_value = st.secrets[key]
    except Exception:
        return ""

    return str(secret_value)


def load_credentials() -> tuple[str, str]:
    client_id = read_config_value("PDF_SERVICES_CLIENT_ID").strip()
    client_secret = read_config_value("PDF_SERVICES_CLIENT_SECRET").strip()
    return client_id, client_secret


def make_safe_sheet_name(name: str, fallback: str) -> str:
    sanitized = "".join(char if char not in r'[]:*?/\\' else "_" for char in name).strip()
    sanitized = sanitized[:31]
    return sanitized or fallback


def parse_cell_value(value: str | None) -> str | int | float:
    if value is None:
        return ""

    text = str(value).strip()
    if not text:
        return ""

    normalized = text.replace(",", "")
    try:
        if normalized.isdigit() or (normalized.startswith("-") and normalized[1:].isdigit()):
            return int(normalized)
        return float(normalized)
    except ValueError:
        return text


def write_excel_cell(sheet, row: int, column: int, value: str | int | float) -> None:
    cell = sheet.cell(row=row, column=column, value=value)
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        cell.number_format = "#,##0"


def is_inside_bbox(word: dict, bbox: tuple[float, float, float, float]) -> bool:
    x0, top, x1, bottom = bbox
    word_x0 = float(word.get("x0", 0.0))
    word_x1 = float(word.get("x1", 0.0))
    word_top = float(word.get("top", 0.0))
    word_bottom = float(word.get("bottom", 0.0))
    return word_x0 >= x0 and word_x1 <= x1 and word_top >= top and word_bottom <= bottom


def group_words_by_line(words: list[dict], tolerance: float = 3.0) -> list[list[dict]]:
    if not words:
        return []

    ordered_words = sorted(words, key=lambda word: (float(word.get("top", 0.0)), float(word.get("x0", 0.0))))
    lines: list[list[dict]] = []
    current_line: list[dict] = []
    current_top: float | None = None

    for word in ordered_words:
        word_top = float(word.get("top", 0.0))
        if not current_line:
            current_line = [word]
            current_top = word_top
            continue

        if current_top is not None and abs(word_top - current_top) <= tolerance:
            current_line.append(word)
            current_top = (current_top + word_top) / 2
            continue

        lines.append(sorted(current_line, key=lambda item: float(item.get("x0", 0.0))))
        current_line = [word]
        current_top = word_top

    if current_line:
        lines.append(sorted(current_line, key=lambda item: float(item.get("x0", 0.0))))

    return lines


def convert_pdf_to_xlsx(uploaded_pdf: bytes) -> bytes:
    workbook = Workbook()
    workbook.remove(workbook.active)

    with pdfplumber.open(io.BytesIO(uploaded_pdf)) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            sheet = workbook.create_sheet(title=make_safe_sheet_name(f"page_{page_index}", fallback=f"page_{page_index}"))
            current_row = 1

            tables = page.find_tables() or []
            table_regions = [table.bbox for table in tables]

            for table in tables:
                extracted_rows = table.extract() or []
                for row_values in extracted_rows:
                    for column_index, cell_value in enumerate(row_values, start=1):
                        write_excel_cell(sheet, current_row, column_index, parse_cell_value(cell_value))
                    current_row += 1
                current_row += 1

            words = page.extract_words(x_tolerance=1, y_tolerance=3, keep_blank_chars=False) or []
            non_table_words = [
                word
                for word in words
                if not any(is_inside_bbox(word, bbox) for bbox in table_regions)
            ]

            for line_words in group_words_by_line(non_table_words):
                for column_index, word in enumerate(line_words, start=1):
                    write_excel_cell(sheet, current_row, column_index, parse_cell_value(word.get("text")))
                current_row += 1

            if current_row == 1:
                sheet.cell(row=1, column=1, value="このページから抽出できる表・文字データは見つかりませんでした。")

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return output.read()


def extract_ocr_lines(image: Image.Image) -> list[dict[str, float | str]]:
    ocr_data = pytesseract.image_to_data(
        image,
        lang="jpn+eng",
        config="--oem 3 --psm 6",
        output_type=Output.DICT,
    )

    grouped_lines: dict[tuple[int, int, int], list[dict[str, float | str]]] = {}
    item_count = len(ocr_data.get("text", []))

    for index in range(item_count):
        text = str(ocr_data["text"][index]).strip()
        confidence_text = str(ocr_data["conf"][index]).strip()
        if not text:
            continue

        try:
            confidence = float(confidence_text)
        except ValueError:
            confidence = -1.0

        if confidence < 0:
            continue

        key = (
            int(ocr_data["block_num"][index]),
            int(ocr_data["par_num"][index]),
            int(ocr_data["line_num"][index]),
        )
        grouped_lines.setdefault(key, []).append(
            {
                "text": text,
                "left": float(ocr_data["left"][index]),
                "top": float(ocr_data["top"][index]),
                "width": float(ocr_data["width"][index]),
                "height": float(ocr_data["height"][index]),
            }
        )

    merged_lines: list[dict[str, float | str]] = []
    for line_items in grouped_lines.values():
        ordered_items = sorted(line_items, key=lambda item: float(item["left"]))
        text = " ".join(str(item["text"]) for item in ordered_items).strip()
        if not text:
            continue

        left = min(float(item["left"]) for item in ordered_items)
        top = min(float(item["top"]) for item in ordered_items)
        height = max(float(item["height"]) for item in ordered_items)
        merged_lines.append(
            {
                "text": text,
                "left": left,
                "top": top,
                "height": height,
            }
        )

    return sorted(merged_lines, key=lambda item: (float(item["top"]), float(item["left"])))


def append_ocr_page_to_doc(document: Document, page_image: Image.Image, dpi: int) -> None:
    lines = extract_ocr_lines(page_image)

    previous_bottom = 0.0
    for line in lines:
        top = float(line["top"])
        height = float(line["height"])
        left = float(line["left"])

        paragraph = document.add_paragraph()
        if previous_bottom > 0:
            vertical_gap_px = max(top - previous_bottom, 0.0)
            paragraph.paragraph_format.space_before = Pt((vertical_gap_px / dpi) * 72)

        paragraph.paragraph_format.left_indent = Inches(left / dpi)
        run = paragraph.add_run(str(line["text"]))
        run.font.size = Pt(max((height / dpi) * 72 * 0.9, 9))

        previous_bottom = top + height

    if not lines:
        document.add_paragraph("OCRで抽出できる文字が見つかりませんでした。")


def convert_pdf_to_docx(uploaded_pdf: bytes) -> bytes:
    dpi = 300
    images = convert_from_bytes(uploaded_pdf, dpi=dpi)
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    for page_index, image in enumerate(images):
        append_ocr_page_to_doc(document, image, dpi=dpi)
        if page_index < len(images) - 1:
            document.add_page_break()

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.read()


def create_pdf_services(client_id: str, client_secret: str) -> PDFServices:
    credentials = ServicePrincipalCredentials(client_id=client_id, client_secret=client_secret)
    return PDFServices(credentials=credentials)


def export_pdf_with_adobe(uploaded_pdf: bytes, output_label: str, client_id: str, client_secret: str) -> bytes:
    pdf_services = create_pdf_services(client_id, client_secret)
    input_asset = pdf_services.upload(input_stream=uploaded_pdf, mime_type=PDFServicesMediaType.PDF)

    export_pdf_params = ExportPDFParams(target_format=OUTPUT_FORMATS[output_label]["target_format"])
    export_pdf_job = ExportPDFJob(input_asset=input_asset, export_pdf_params=export_pdf_params)

    location = pdf_services.submit(export_pdf_job)
    pdf_services_response = pdf_services.get_job_result(location, ExportPDFResult)
    result_asset: CloudAsset = pdf_services_response.get_result().get_asset()
    stream_asset: StreamAsset = pdf_services.get_content(result_asset)
    return stream_asset.get_input_stream()


def run_conversion(uploaded_file: st.runtime.uploaded_file_manager.UploadedFile, output_label: str, client_id: str, client_secret: str) -> tuple[bytes, str, str]:
    if output_label == "Word (.docx)":
        output_bytes = convert_pdf_to_docx(uploaded_file.getvalue())
    elif output_label == "Excel (.xlsx)":
        output_bytes = convert_pdf_to_xlsx(uploaded_file.getvalue())
    else:
        output_bytes = export_pdf_with_adobe(uploaded_file.getvalue(), output_label, client_id, client_secret)

    meta = OUTPUT_FORMATS[output_label]
    output_name = f"{Path(uploaded_file.name).stem}.{meta['extension']}"
    return output_bytes, output_name, meta["mime"]


def main() -> None:
    set_page_style()
    client_id, client_secret = load_credentials()

    st.markdown(
        """
        <div class="hero">
            <h1>PDF 変換アプリ</h1>
            <p>PDF を Word / Excel / PowerPoint の編集可能形式へ変換します。</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    uploaded_file = st.file_uploader("PDFファイルをアップロード", type=["pdf"])
    output_label = st.selectbox("変換先の形式", list(OUTPUT_FORMATS.keys()))

    credentials_ready = bool(client_id and client_secret)
    adobe_required = output_label == "PowerPoint (.pptx)"
    if adobe_required and not credentials_ready:
        st.warning("管理者にお問い合わせください（Secrets未設定）")

    convert_clicked = st.button(
        "変換開始",
        use_container_width=True,
        disabled=uploaded_file is None or (adobe_required and not credentials_ready),
    )

    if convert_clicked and uploaded_file is not None and (credentials_ready or not adobe_required):
        try:
            with st.spinner("変換中です。ファイルサイズや内容によって時間がかかる場合があります。"):
                output_bytes, output_name, mime_type = run_conversion(uploaded_file, output_label, client_id, client_secret)

            st.session_state["converted_file"] = {
                "bytes": output_bytes,
                "name": output_name,
                "mime": mime_type,
            }
            st.success("変換が完了しました。")
        except (ServiceApiException, ServiceUsageException, SdkException) as error:
            st.session_state.pop("converted_file", None)
            st.error(f"Adobe PDF Services で変換に失敗しました: {error}")
        except Exception as error:
            st.session_state.pop("converted_file", None)
            st.error(f"変換に失敗しました: {error}")

    converted_file = st.session_state.get("converted_file")
    if converted_file:
        st.download_button(
            label="変換ファイルをダウンロード",
            data=converted_file["bytes"],
            file_name=converted_file["name"],
            mime=converted_file["mime"],
            use_container_width=True,
        )


if __name__ == "__main__":
    main()